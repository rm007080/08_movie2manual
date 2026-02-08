"""
Wordドキュメント生成モジュール

python-docxを使用してマニュアルのWordファイルを生成する。
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict, Optional
from pathlib import Path
import io
import logging

logger = logging.getLogger(__name__)


class DocGenerator:
    """Wordマニュアル生成クラス"""

    def __init__(self, project_name: str = "操作マニュアル"):
        """初期化

        Args:
            project_name: ドキュメントのタイトル（プロジェクト名）
        """
        self.doc = Document()
        self.project_name = project_name
        self._setup_styles()
        self._setup_page()

    def _setup_styles(self):
        """ドキュメントのスタイル設定"""
        # デフォルトフォントを設定（日本語対応）
        self.doc.styles['Normal'].font.name = 'MS Mincho'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')
        self.doc.styles['Normal'].font.size = Pt(11)

        # 段落設定
        paragraph_format = self.doc.styles['Normal'].paragraph_format
        paragraph_format.line_spacing = 1.5
        paragraph_format.space_after = Pt(6)

        # Heading 2 スタイル（ステップ見出し用 - 目次対応）
        h2 = self.doc.styles['Heading 2']
        h2.font.name = 'MS Mincho'
        h2._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 0, 0)
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after = Pt(6)

    def _setup_page(self):
        """ページ設定（余白、ヘッダー、フッター）"""
        section = self.doc.sections[0]

        # 余白設定
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

        # ヘッダー（プロジェクト名・右寄せ）
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header_para.add_run(self.project_name)
        run.font.size = Pt(9)
        run.font.name = 'MS Mincho'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')
        run.font.color.rgb = RGBColor(128, 128, 128)

        # フッター（ページ番号・中央揃え）
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        prefix = footer_para.add_run("- ")
        prefix.font.size = Pt(9)
        prefix.font.color.rgb = RGBColor(128, 128, 128)

        self._add_field(footer_para, ' PAGE ')

        sep = footer_para.add_run(" / ")
        sep.font.size = Pt(9)
        sep.font.color.rgb = RGBColor(128, 128, 128)

        self._add_field(footer_para, ' NUMPAGES ')

        suffix = footer_para.add_run(" -")
        suffix.font.size = Pt(9)
        suffix.font.color.rgb = RGBColor(128, 128, 128)

    def _add_field(self, paragraph, field_code):
        """Wordフィールドコードを段落に追加

        Args:
            paragraph: 対象段落
            field_code: フィールドコード（例: ' PAGE ', ' NUMPAGES '）
        """
        run = paragraph.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar_begin)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = field_code
        run._element.append(instrText)

        fldChar_separate = OxmlElement('w:fldChar')
        fldChar_separate.set(qn('w:fldCharType'), 'separate')
        run._element.append(fldChar_separate)

        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar_end)

        return run

    def add_title(self, title: str):
        """ドキュメントのタイトルを追加

        Args:
            title: タイトルテキスト
        """
        # タイトル用の段落を追加
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title_run = title_para.add_run(title)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.name = 'MS Mincho'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')

        # タイトル後のスペース
        title_para.paragraph_format.space_after = Pt(20)

    def add_toc(self):
        """目次を追加（Heading 2 のステップ見出しを収集）

        dirty フラグにより、Wordで開いた際にページレイアウト確定後に
        フィールド更新が促される。
        """
        toc_heading = self.doc.add_paragraph()
        toc_run = toc_heading.add_run("目次")
        toc_run.font.size = Pt(14)
        toc_run.font.bold = True
        toc_run.font.name = 'MS Mincho'
        toc_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')
        toc_heading.paragraph_format.space_after = Pt(6)

        # TOCフィールド（dirty フラグ + 分割 run 方式）
        toc_para = self.doc.add_paragraph()

        # Begin（dirty=true でレイアウト確定後に再計算を強制）
        run_begin = toc_para.add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        fldChar_begin.set(qn('w:dirty'), 'true')
        run_begin._element.append(fldChar_begin)

        # Instruction
        run_instr = toc_para.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = r' TOC \o "2-2" \h \z \u '
        run_instr._element.append(instrText)

        # Separate
        run_sep = toc_para.add_run()
        fldChar_sep = OxmlElement('w:fldChar')
        fldChar_sep.set(qn('w:fldCharType'), 'separate')
        run_sep._element.append(fldChar_sep)

        # End
        run_end = toc_para.add_run()
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run_end._element.append(fldChar_end)

        self.doc.add_page_break()

    def add_step_header(self, step_num: int, title: str):
        """ステップの見出しを追加（Heading 2 スタイル - 目次対応）

        Args:
            step_num: ステップ番号
            title: ステップのタイトル
        """
        header_text = f"ステップ{step_num}: {title}"
        self.doc.add_heading(header_text, level=2)

    def add_step_image(self, image_path: str, max_width: float = 5.0):
        """ステップの画像を追加（中央揃え）

        Args:
            image_path: 画像ファイルのパス
            max_width: 画像の最大幅（インチ）、デフォルト5インチ（約12.7cm）
        """
        img_path = Path(image_path)

        if not img_path.exists():
            logger.warning(f"Image file not found: {image_path}")
            # 画像がない場合はプレースホルダーテキストを追加
            placeholder_para = self.doc.add_paragraph()
            placeholder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            placeholder_run = placeholder_para.add_run(f"[画像: {img_path.name}]")
            placeholder_run.font.italic = True
            placeholder_run.font.color.rgb = RGBColor(150, 150, 150)
            return

        try:
            # 画像を中央揃えで追加
            img_para = self.doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 画像を追加（幅を指定）
            run = img_para.add_run()
            run.add_picture(str(img_path.absolute()), width=Inches(max_width))

            # 画像周りのスペース
            img_para.paragraph_format.space_before = Pt(6)
            img_para.paragraph_format.space_after = Pt(6)

        except Exception as e:
            logger.error(f"Error adding image {image_path}: {e}")
            # エラー時はプレースホルダーを追加
            placeholder_para = self.doc.add_paragraph()
            placeholder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            placeholder_run = placeholder_para.add_run(f"[画像読み込みエラー: {img_path.name}]")
            placeholder_run.font.italic = True
            placeholder_run.font.color.rgb = RGBColor(200, 50, 50)

    def add_step_description(self, description: str):
        """ステップの説明文を追加

        Args:
            description: 説明文
        """
        desc_para = self.doc.add_paragraph()

        desc_run = desc_para.add_run(description)
        desc_run.font.size = Pt(11)
        desc_run.font.name = 'MS Mincho'
        desc_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')

        desc_para.paragraph_format.line_spacing = 1.5
        desc_para.paragraph_format.space_after = Pt(12)

    def add_step(self, step_num: int, title: str, image_path: str, description: str):
        """1ステップ分のコンテンツを追加

        Args:
            step_num: ステップ番号
            title: ステップのタイトル
            image_path: 画像ファイルのパス
            description: 説明文
        """
        self.add_step_header(step_num, title)
        self.add_step_image(image_path)
        self.add_step_description(description)

    def save(self, output_path: str):
        """ドキュメントを保存

        Args:
            output_path: 出力ファイルパス

        Returns:
            保存したファイルの絶対パス
        """
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)

        self.doc.save(str(output.absolute()))
        logger.info(f"Document saved to: {output.absolute()}")
        return str(output.absolute())

    def get_bytes(self) -> bytes:
        """ドキュメントをバイト列として取得（Streamlitのダウンロード用）

        Returns:
            ドキュメントのバイト列
        """
        # 一時的なバッファにドキュメントを保存
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def add_page_break(self):
        """改ページを追加"""
        self.doc.add_page_break()

    def add_section(self, title: str):
        """セクション見出しを追加

        Args:
            title: セクションタイトル
        """
        section_para = self.doc.add_paragraph()

        section_run = section_para.add_run(title)
        section_run.font.size = Pt(16)
        section_run.font.bold = True
        section_run.font.name = 'MS Mincho'
        section_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')

        section_para.paragraph_format.space_before = Pt(18)
        section_para.paragraph_format.space_after = Pt(12)


def create_word_manual(steps_data: List[Dict], output_path: str,
                       project_name: str = "操作マニュアル") -> str:
    """マニュアルWordファイルを作成する便利関数

    Args:
        steps_data: ステップデータのリスト
            [{
                "id": 1,
                "title": "メニューを開く",
                "description": "「インストールされているアプリ」をクリックします。",
                "image_path": "path/to/image.jpg"
            }, ...]
        output_path: 出力ファイルパス
        project_name: プロジェクト名（ドキュメントタイトル）

    Returns:
        出力ファイルの絶対パス
    """
    gen = DocGenerator(project_name)
    gen.add_title(project_name)

    for step in steps_data:
        step_num = step.get("id", step.get("step_num", 0))
        title = step.get("title", "")
        image_path = step.get("image_path", "")
        description = step.get("description", "")

        gen.add_step(step_num, title, image_path, description)

    return gen.save(output_path)


def create_word_manual_from_paths(image_paths: List[str], descriptions: List[str],
                                  output_path: str, project_name: str = "操作マニュアル") -> str:
    """画像パスと説明文のリストからマニュアルを作成する簡易関数

    Args:
        image_paths: 画像ファイルパスのリスト
        descriptions: 説明文のリスト（image_pathsと同じ長さ）
        output_path: 出力ファイルパス
        project_name: プロジェクト名

    Returns:
        出力ファイルの絶対パス
    """
    if len(image_paths) != len(descriptions):
        raise ValueError("image_pathsとdescriptionsの長さが一致しません")

    steps_data = [
        {
            "id": i + 1,
            "title": f"手順 {i + 1}",
            "image_path": img_path,
            "description": desc
        }
        for i, (img_path, desc) in enumerate(zip(image_paths, descriptions))
    ]

    return create_word_manual(steps_data, output_path, project_name)
